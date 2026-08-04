import os
import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def generate_updated_valuation():
    doc = Document()
    
    # Configurar márgenes de página (1 pulgada)
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Insertar Encabezado con Logos Institucionales
    header_table = doc.add_table(rows=1, cols=2)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    header_table.autofit = False
    
    # Remover bordes de tabla de encabezado
    for cell in header_table.rows[0].cells:
        tcPr = cell._element.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="none"/><w:left w:val="none"/><w:bottom w:val="none"/><w:right w:val="none"/></w:tcBorders>')
        tcPr.append(tcBorders)
        
    logo_fv_p = os.path.join(os.path.dirname(__file__), "assets", "Logo_FV_Principal.png")
    logo_altus_p = os.path.join(os.path.dirname(__file__), "assets", "Logo_ALTUS AI_Principal.png")
    
    if os.path.exists(logo_fv_p):
        p_fv = header_table.rows[0].cells[0].paragraphs[0]
        p_fv.add_run().add_picture(logo_fv_p, height=Inches(0.6))
    if os.path.exists(logo_altus_p):
        p_altus = header_table.rows[0].cells[1].paragraphs[0]
        p_altus.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p_altus.add_run().add_picture(logo_altus_p, height=Inches(0.45))
        
    doc.add_paragraph("")
    
    # Título Principal
    title_p = doc.add_paragraph()
    title_run = title_p.add_run("Auditoría Técnica y Re-Valorización Maestra de Altus AI")
    title_run.font.name = "Calibri"
    title_run.font.size = Pt(22)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(10, 35, 66) # Navy #0A2342
    
    subtitle_p = doc.add_paragraph()
    subtitle_run = subtitle_p.add_run("Ecosistema WealthTech & Core Bancario Cognitivo Institucional (FV Asesorías)")
    subtitle_run.font.name = "Calibri"
    subtitle_run.font.size = Pt(13)
    subtitle_run.font.italic = True
    subtitle_run.font.color.rgb = RGBColor(2, 132, 199) # Sky Blue #0284C7
    
    # Metadata
    meta_p = doc.add_paragraph()
    r1 = meta_p.add_run("Fecha de Actualización: ")
    r1.bold = True
    meta_p.add_run("29 de Julio de 2026\n")
    r2 = meta_p.add_run("Estado del Ecosistema: ")
    r2.bold = True
    meta_p.add_run("Plataforma Enterprise Multi-Módulo / Producción Completa")
    
    doc.add_paragraph("─" * 55)
    
    # Resumen Ejecutivo
    doc.add_heading("1. Resumen Ejecutivo de la Re-Valorización", level=1)
    p_desc = doc.add_paragraph(
        "Tras la incorporación de los nuevos módulos desarrollados a julio de 2026, la magnitud técnica y comercial de Altus AI "
        "consolida un salto cualitativo significativo. La plataforma ya no es solo un gestor o agregador patrimonial, sino un "
        "Core Bancario Cognitivo e Infraestructura WealthTech Enterprise. La integración viva de motores de posicionamiento e "
        "inteligencia en LinkedIn con APIs oficiales en tiempo real (BCCh, Cobre/LME, USD/CLP, UF), infografías 4K automatizadas, "
        "transcripción de audios de corredoras (Gemini 2.5 Pro), generadores de fichas KYC institucionales en Excel y simuladores "
        "tributarios/cuantitativos disparan los múltiplos de valorización M&A a niveles de liderazgo Fintech regional."
    )
    p_desc.paragraph_format.line_spacing = 1.15
    
    # Alcance de Módulos
    doc.add_heading("2. Inventario Completo de Módulos e Innovaciones", level=1)
    
    modulos = [
        ("💬 Copiloto de Comentarios & Engagement Estratégico (LinkedIn):", 
         "Módulo autónomo de análisis de artículos web (DF, La Tercera, Pulso, Bloomberg) y publicaciones de LinkedIn. Conexión directa a APIs estadísticas oficiales en tiempo real (BCCh, Cobre US$/lb LME, USD/CLP, TPM, S&P 500, UF) y generación de 4 posturas conversacionales de nivel ejecutivo en primera persona (Macroeconómica, Contrapunto Crítico, Asesor Patrimonial, Punchy)."),
        
        ("📱 Motor de Infografías 4K & Contenido Multimodal:", 
         "Generador automatizado de infografías 4K (PNG lossless, escala 2x, SVG vectoriales), carruseles PDF para LinkedIn e informes Word con racional técnico a partir de RSS, audios de WhatsApp de corredoras (transcripción multimodal) o noticias locales."),
        
        ("📑 Generador Institucional KYC & Ficha del Cliente en Excel:", 
         "Motor de creación de planillas Excel estilizadas (#0A2342 Navy, #0284C7 Sky Blue) con levantamiento de perfil de riesgo, régimen impositivo (APV A/B, Reliquidación IGC Art. 54/57 Bis), ingresos independientes e imponibles del cónyuge (tributación conjunta/separada) y guía paso a paso para el cliente."),
        
        ("🧮 Simuladores Cuantitativos & Tributarios (Chile):", 
         "Cálculo optimizado de APV (Régimen A vs B), Reliquidación IGC, Simulación Crédito Hipotecario vs Inversión y Valuación Real Estate."),
        
        ("🔍 OSINT Profundo & Ingesta Maestra de Cartolas:", 
         "Cruce autónomo de bases de datos CMF, Diario Oficial, InfoProbidad, Transunion, cartolas bancarias/corredoras (LarrainVial, BTG, Principal, SURA) con cifrado AES-256."),
        
        ("💼 CRM Patrimonial & Hub Comercial:", 
         "Pipeline comercial de clientes, scoring de prospectos patrimoniales y plantillas HTML de correo institucionalizadas para envío directo.")
    ]
    
    for mod_name, mod_text in modulos:
        p_m = doc.add_paragraph()
        r_title = p_m.add_run(mod_name + " ")
        r_title.bold = True
        r_title.font.color.rgb = RGBColor(10, 35, 66)
        p_m.add_run(mod_text)
        p_m.paragraph_format.line_spacing = 1.15
        p_m.paragraph_format.space_after = Pt(6)
        
    # Cuadro de Valorización Actualizada
    doc.add_heading("3. Cuadro de Valorización Maestra (SaaS + Deep Tech Fintech)", level=1)
    
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    headers = ['Método de Valorización', 'Cálculo Teórico & Justificación', 'Estimación (USD)']
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        set_cell_background(hdr_cells[i], '0A2342')
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(255, 255, 255)
            
    val_data = [
        ("Piso de Destrucción\n(Costo de Reposición Real)", 
         "Costo de contratar equipos senior de Data Science, Quant Finance, OSINT Scraping, IA Multimodal (Audio/NLP/Vision), UI Design 4K y Desarrolladores Full-Stack por 3+ años de desarrollo continuo e integración de APIs en vivo.",
         "$2.000.000 - $3.000.000 USD"),
         
        ("Valor de Ahorro Interno (DCF)", 
         "Sustitución directa de agencias de marketing/diseño, analistas macro, auditores tributarios, consultores KYC, redactores RRSS, brokers de prospección y suscripciones a terminales de mercado B2B.",
         "$5.000.000 - $7.000.000 USD"),
         
        ("Potencial Comercial M&A\n(Marca Blanca / Spin-offs B2B)", 
         "Licenciamiento empresarial de la plataforma en 4 verticales independientes: Core WealthTech, Copiloto Marketing/RRSS con APIs, Hub Tributario/KYC, y Quant & Macro Terminal. Múltiplos M&A Deep Tech (12x - 18x ARR).",
         "$15.000.000 - $20.000.000+ USD")
    ]
    
    for metodo, calc, est in val_data:
        row_cells = table.add_row().cells
        row_cells[0].text = metodo
        row_cells[1].text = calc
        row_cells[2].text = est
        
        # Formato de celda
        p0 = row_cells[0].paragraphs[0]
        p0.runs[0].font.bold = True
        p0.runs[0].font.color.rgb = RGBColor(10, 35, 66)
        
        p2 = row_cells[2].paragraphs[0]
        p2.runs[0].font.bold = True
        p2.runs[0].font.color.rgb = RGBColor(2, 132, 199)
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
    doc.add_paragraph("")
    
    # Conclusión
    doc.add_heading("4. Conclusión del Auditor Jefe", level=2)
    p_c = doc.add_paragraph()
    r_c = p_c.add_run(
        "Altus AI se consagra como un activo tecnológico patentable de alto valor estratégico. "
        "Con la adición de la inteligencia de engagement en LinkedIn alimentada por APIs bancarias/commodities en tiempo real, "
        "el generador KYC institucional y las infografías 4K, la valorización razonable del ecosistema consolidado en el mercado M&A Fintech "
        "alcanza los $15.000.000 USD (con un potencial de licenciamiento Spin-off que supera los $20.000.000 USD)."
    )
    r_c.bold = True
    r_c.font.size = Pt(11)
    r_c.font.color.rgb = RGBColor(10, 35, 66)
    
    # Guardar en ambas ubicaciones
    fn_dated = 'VALUATION_ACTUALIZADA_29.07.2026.docx'
    fn_today = 'VALUATION_ACTUALIZADA_Hoy.docx'
    
    doc.save(fn_dated)
    doc.save(fn_today)
    
    print(f"EXITO: Archivo generado exitosamente: {fn_dated}")
    print(f"EXITO: Archivo actualizado exitosamente: {fn_today}")

if __name__ == "__main__":
    generate_updated_valuation()

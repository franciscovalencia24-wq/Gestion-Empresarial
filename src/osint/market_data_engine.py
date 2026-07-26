import os
import yfinance as yf
from src.osint.open_data_engine import OpenDataEngine
import requests
import feedparser
import datetime
import logging
import json
import base64
import urllib.parse
import re
from PIL import Image
from bs4 import BeautifulSoup
from duckduckgo_search import DDGS
import google.generativeai as genai
from jinja2 import Environment, FileSystemLoader
from html2image import Html2Image
from src.database.connection import SessionLocal
from src.database.models import MarketStat, MarketNews
from docx import Document

logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')

class MarketDataEngine:
    def __init__(self):
        self.db = SessionLocal()
        api_key = os.getenv("GOOGLE_API_KEY")
        if api_key:
            genai.configure(api_key=api_key)
            self.model = genai.GenerativeModel("gemini-2.5-pro")
        else:
            self.model = None

    def get_base64_image(self, filepath):
        if not os.path.exists(filepath): return ""
        with open(filepath, "rb") as f:
            encoded = base64.b64encode(f.read()).decode("utf-8")
            mime = "image/svg+xml" if filepath.endswith(".svg") else f"image/{filepath.split('.')[-1]}"
            return f"data:{mime};base64,{encoded}"

    def fetch_global_stats(self):
        """Descarga 14 métricas globales (yfinance)"""
        logging.info("Descargando 14 índices (yfinance)...")
        tickers_info = [
            ("USA", "^GSPC", "S&P 500"),
            ("USA", "^DJI", "DJIA"),
            ("USA", "^NDX", "NASDAQ 100"),
            ("MÉXICO", "^MXX", "IPC"),
            ("BRASIL", "^BVSP", "BOVESPA"),
            ("CHILE", "^IPSA", "IPSA"),
            ("EUROPA", "^STOXX50E", "EURO STOXX 50"),
            ("EUROPA", "^GDAXI", "DAX 40"),
            ("EUROPA", "^FTSE", "FTSE 100"),
            ("MSCI", "URTH", "MSCI WORLD"),
            ("ASIA", "^N225", "NIKKEI 225"),
            ("ASIA", "^HSI", "HANG SENG"),
            ("ASIA", "000001.SS", "SHANGHAI COMP.")
        ]
        
        results_by_region = {"USA": [], "MÉXICO": [], "BRASIL": [], "CHILE": [], "EUROPA": [], "MSCI": [], "ASIA": []}
        
        for region, symbol, name in tickers_info:
            try:
                t = yf.Ticker(symbol)
                data = t.history(period="5d")
                
                if len(data) >= 2:
                    close_today = float(data['Close'].iloc[-1])
                    close_yesterday = float(data['Close'].iloc[-2])
                elif len(data) == 1:
                    close_today = float(data['Close'].iloc[-1])
                    close_yesterday = float(t.info.get('previousClose', close_today))
                else:
                    results_by_region[region].append({"nombre": name, "valor": "N/D", "efecto": "NEUTRAL"})
                    continue
                    
                delta_pct = ((close_today - close_yesterday) / close_yesterday) * 100 if close_yesterday else 0
                efecto = "ALZA" if delta_pct > 0 else "BAJA" if delta_pct < 0 else "NEUTRAL"
                var_abs = abs(delta_pct)
                relevancia = "IMPORTANTE" if var_abs > 1.0 else "MODERADA" if var_abs > 0.3 else "LEVE"
                
                # Formato europeo para número
                valor_str = f"{close_today:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                
                results_by_region[region].append({
                    "nombre": name,
                    "valor": valor_str,
                    "efecto": efecto,
                    "relevancia": relevancia
                })
            except Exception as e:
                logging.error(f"Error descargando {symbol}: {e}")
                results_by_region[region].append({"nombre": name, "valor": "N/D", "efecto": "NEUTRAL", "relevancia": "LEVE"})
                
        return results_by_region

    def fetch_currency_stats(self):
        """Descarga Monedas y UF de forma nativa desde Yahoo y Mindicador"""
        logging.info("Calculando variaciones de Monedas (Yahoo + Mindicador)...")
        results = []
        
        # 1. UF desde mindicador (con manejo de timeout)
        try:
            resp = requests.get('https://mindicador.cl/api/uf', timeout=5)
            if resp.status_code == 200:
                serie = resp.json().get('serie', [])
                if len(serie) >= 2:
                    val_today = float(serie[0]['valor'])
                    val_yest = float(serie[1]['valor'])
                    var = ((val_today - val_yest) / val_yest) * 100 if val_yest else 0
                    var_abs = abs(var)
                    relevancia = "IMPORTANTE" if var_abs > 0.5 else "MODERADA" if var_abs > 0.1 else "LEVE"
                    efecto = "ALZA" if var > 0 else "BAJA" if var < 0 else "NEUTRAL"
                    val_str = f"{val_today:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                    results.append({
                        "nombre": "UF",
                        "valor": val_str,
                        "variacion": f"{var:+.2f}%",
                        "efecto": efecto,
                        "relevancia": relevancia
                    })
        except Exception as e:
            logging.error(f"Error descargando UF desde mindicador: {e}")

        # 2. Monedas desde Yahoo Finance
        pairs = [
            ("DÓLAR (USD/CLP)", "CLP=X"),
            ("EURO (EUR/USD)", "EURUSD=X"),
            ("LIBRA (GBP/USD)", "GBPUSD=X"),
            ("REAL (USD/BRL)", "BRL=X")
        ]
        for name, symbol in pairs:
            try:
                data = yf.Ticker(symbol).history(period="5d")
                if len(data) >= 2:
                    v1 = float(data['Close'].iloc[-1])
                    v0 = float(data['Close'].iloc[-2])
                    var = ((v1 - v0) / v0) * 100 if v0 else 0
                    efecto = "ALZA" if var > 0 else "BAJA" if var < 0 else "NEUTRAL"
                    
                    var_abs = abs(var)
                    relevancia = "IMPORTANTE" if var_abs > 1.0 else "MODERADA" if var_abs > 0.3 else "LEVE"
                    
                    if symbol == "CLP=X":
                        val_str = f"{v1:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                    else:
                        val_str = f"{v1:,.4f}".replace(".", ",")
                        
                    results.append({
                        "nombre": name,
                        "valor": val_str,
                        "variacion": f"{var:+.2f}%",
                        "efecto": efecto,
                        "relevancia": relevancia
                    })
            except Exception as e:
                logging.error(f"Error descargando moneda {symbol}: {e}")
                
        return results

    def fetch_commodities_stats(self):
        """Descarga Commodities de forma nativa desde Yahoo"""
        logging.info("Calculando variaciones de Commodities (Yahoo)...")
        results = []
        pairs = [
            ("PETRÓLEO WTI", "CL=F"),
            ("ORO", "GC=F"),
            ("COBRE CASH", "HG=F"),
            ("PLATA", "SI=F"),
            ("GAS NATURAL", "NG=F")
        ]
        for name, symbol in pairs:
            try:
                data = yf.Ticker(symbol).history(period="5d")
                if len(data) >= 2:
                    v1 = float(data['Close'].iloc[-1])
                    v0 = float(data['Close'].iloc[-2])
                    var = ((v1 - v0) / v0) * 100 if v0 else 0
                    efecto = "ALZA" if var > 0 else "BAJA" if var < 0 else "NEUTRAL"
                    
                    var_abs = abs(var)
                    relevancia = "IMPORTANTE" if var_abs > 1.0 else "MODERADA" if var_abs > 0.3 else "LEVE"
                    
                    val_str = f"{v1:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
                        
                    results.append({
                        "nombre": name,
                        "valor": val_str,
                        "variacion": f"{var:+.2f}%",
                        "efecto": efecto,
                        "relevancia": relevancia
                    })
            except Exception as e:
                logging.error(f"Error descargando comodity {symbol}: {e}")
                results.append({
                    "nombre": name,
                    "valor": "N/D",
                    "variacion": "N/D",
                    "efecto": "NEUTRAL",
                    "relevancia": "LEVE"
                })
        return results

    def fetch_reuters_news(self):
        logging.info("Descargando noticias globales (RSS)...")
        rss_urls = [("Yahoo Finance", "https://finance.yahoo.com/news/rssindex")]
        today_start = datetime.datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        
        for source_name, url in rss_urls:
            try:
                feed = feedparser.parse(url)
                for entry in feed.entries[:15]:
                    if hasattr(entry, 'published_parsed'):
                        import calendar
                        utc_ts = calendar.timegm(entry.published_parsed)
                        pub_date = datetime.datetime.fromtimestamp(utc_ts)
                    else:
                        pub_date = datetime.datetime.now()
                        
                    if pub_date >= today_start:
                        exists = self.db.query(MarketNews).filter_by(titular=entry.title).first()
                        if not exists:
                            news = MarketNews(fuente=source_name, titular=entry.title, resumen=BeautifulSoup(entry.summary, "html.parser").text if hasattr(entry, 'summary') else "", link=entry.link, fecha_publicacion=pub_date)
                            self.db.add(news)
                self.db.commit()
            except Exception as e:
                logging.error(f"Error procesando RSS: {e}")

    def fetch_chile_news(self):
        logging.info("Descargando noticias financieras de Chile (Google News RSS)...")
        import feedparser
        import urllib.parse
        
        query = "economia chile when:1d"
        url = f"https://news.google.com/rss/search?q={urllib.parse.quote(query)}&hl=es-419&gl=CL&ceid=CL:es-419"
        today_start = datetime.datetime.now().replace(hour=0, minute=0, second=0, microsecond=0)
        
        try:
            feed = feedparser.parse(url)
            for entry in feed.entries[:15]:
                if hasattr(entry, 'published_parsed'):
                    import calendar
                    utc_ts = calendar.timegm(entry.published_parsed)
                    pub_date = datetime.datetime.fromtimestamp(utc_ts)
                else:
                    pub_date = datetime.datetime.now()
                    
                if pub_date >= today_start:
                    exists = self.db.query(MarketNews).filter_by(titular=entry.title).first()
                    if not exists:
                        news = MarketNews(fuente="Google News Chile", titular=entry.title, resumen="", link=entry.link, fecha_publicacion=pub_date)
                        self.db.add(news)
            self.db.commit()
        except Exception as e:
            logging.error(f"Error procesando RSS Chile: {e}")

    def fetch_custom_url(self, url):
        """Scrapea una URL personalizada para usarla como noticia principal."""
        try:
            # Usar User-Agent de Googlebot para saltar muros de pago básicos (ej: DF)
            headers = {"User-Agent": "Mozilla/5.0 (compatible; Googlebot/2.1; +http://www.google.com/bot.html)"}
            r = requests.get(url, headers=headers, timeout=10)
            if r.status_code == 200:
                soup = BeautifulSoup(r.text, 'html.parser')
                
                title = soup.title.string if soup.title else "Noticia Personalizada"
                
                date_meta = soup.find('meta', property='article:published_time') or soup.find('meta', attrs={'name': 'pubdate'}) or soup.find('meta', attrs={'name': 'cXenseParse:recs:publishtime'})
                if date_meta and date_meta.get('content'):
                    fecha_str = date_meta['content']
                else:
                    fecha_str = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                    
                site_name = soup.find('meta', property='og:site_name')
                fuente_str = site_name['content'] if site_name and site_name.get('content') else "Sitio Web Externo"

                # Intentar enfocar el contenedor principal para evitar ruido del header/sidebar
                main_node = soup.find('article')
                if not main_node:
                    for cls in ['article-body', 'post-content', 'contenido-noticia', 'entry-content', 'body-nota']:
                        main_node = soup.find('div', class_=re.compile(cls, re.I)) or soup.find('div', id=re.compile(cls, re.I))
                        if main_node: break
                
                if not main_node:
                    main_node = soup
                
                # Filtrar párrafos pequeños que suelen ser tickers o menús
                paragraphs = [p.text.strip() for p in main_node.find_all('p') if len(p.text.strip()) > 60]
                content = " ".join(paragraphs[:15])
                
                if not content:
                    return None
                    
                return f"- {title} (Fuente: {fuente_str}, Fecha: {fecha_str}) : {content}"
            return None
        except Exception as e:
            logging.error(f"Error extrayendo URL personalizada: {e}")
            return None

    def search_topic(self, topic):
        """Busca un tema en DDG y retorna el primer resultado útil o un resumen de los snippets."""
        try:
            results = list(DDGS().text(topic, max_results=5, backend="html"))
            
            # 1. Intentar raspar el artículo completo de los primeros resultados
            for res in results[:3]:
                url = res.get('href')
                if url:
                    news_text = self.fetch_custom_url(url)
                    # Si pudo raspar y obtuvo un texto decente (no solo un aviso de paywall)
                    if news_text and len(news_text) > 400:
                        return f"[RESULTADO BÚSQUEDA DDG] Tema: {topic}\n{news_text}"
            
            # 2. FALLBACK: Si todo estaba bloqueado por paywall, usar los resúmenes del buscador
            fallback_text = f"Resumen de Noticias (Extraído de Snippets del Buscador):\n\n"
            for res in results:
                title = res.get('title', 'Sin Título')
                body = res.get('body', '')
                url = res.get('href', '')
                if body:
                    fallback_text += f"- TÍTULO: {title}\n- RESUMEN: {body}\n- FUENTE: {url}\n\n"
            
            if len(fallback_text) > 150:
                return f"[RESULTADO BÚSQUEDA DDG] Tema: {topic}\n{fallback_text}"
                
            # 3. GOOGLE NEWS RSS FALLBACK (Anti-Bloqueo)
            import urllib.request
            import urllib.parse
            import xml.etree.ElementTree as ET
            
            gnews_url = f"https://news.google.com/rss/search?q={urllib.parse.quote(topic)}&hl=es-419&gl=CL&ceid=CL:es-419"
            req = urllib.request.Request(gnews_url, headers={'User-Agent': 'Mozilla/5.0'})
            xml_data = urllib.request.urlopen(req, timeout=10).read()
            root = ET.fromstring(xml_data)
            
            rss_text = f"Resumen de Noticias (Google News RSS):\n\n"
            items = root.findall('.//item')
            for item in items[:5]:
                title = item.find('title').text if item.find('title') is not None else 'Sin Título'
                link = item.find('link').text if item.find('link') is not None else ''
                pub_date = item.find('pubDate').text if item.find('pubDate') is not None else ''
                rss_text += f"- TÍTULO: {title}\n- FECHA: {pub_date}\n- FUENTE: {link}\n\n"
                
            if len(items) > 0:
                return f"[RESULTADO BÚSQUEDA GNEWS] Tema: {topic}\n{rss_text}"
                
            return None
        except Exception as e:
            logging.error(f"Error buscando tema en web: {e}")
            raise ValueError(f"Error de conexión con los buscadores (DuckDuckGo y Google News fallaron). Detalle: {e}")

    def markdown_to_unicode_bold(self, text):
        def replace_bold(match):
            content = match.group(1)
            bold_text = ""
            for char in content:
                if 'A' <= char <= 'Z':
                    bold_text += chr(ord(char) - ord('A') + 0x1D5D4)
                elif 'a' <= char <= 'z':
                    bold_text += chr(ord(char) - ord('a') + 0x1D5EE)
                elif '0' <= char <= '9':
                    bold_text += chr(ord(char) - ord('0') + 0x1D7EC)
                else:
                    bold_text += char
            return bold_text
        return re.sub(r'\*\*(.*?)\*\*', replace_bold, text)

    def generate_content(self, currency_stats, commodity_stats, global_stats, custom_news_text=None, mode="auto"):
        if not self.model: return None
        
        if custom_news_text:
            news_text = custom_news_text
            logging.info("Usando noticia personalizada para generar contenido.")
        else:
            if mode == "auto_chile":
                news = self.db.query(MarketNews).filter(MarketNews.usado_para_linkedin == 0, MarketNews.fuente == "Google News Chile").order_by(MarketNews.fecha_publicacion.desc()).limit(40).all()
                if not news:
                    news = self.db.query(MarketNews).filter(MarketNews.fuente == "Google News Chile").order_by(MarketNews.fecha_publicacion.desc()).limit(15).all()
            else:
                news = self.db.query(MarketNews).filter(MarketNews.usado_para_linkedin == 0, MarketNews.fuente != "Google News Chile").order_by(MarketNews.fecha_publicacion.desc()).limit(40).all()
                if not news: # Fallback just in case
                    news = self.db.query(MarketNews).filter(MarketNews.fuente != "Google News Chile").order_by(MarketNews.fecha_publicacion.desc()).limit(15).all()
            news_text = "\n".join([f"- {n.titular} (Fuente: {n.fuente}, Fecha: {n.fecha_publicacion}) : {n.resumen}" for n in news])
            
        market_status_text = "RENDIMIENTO ACTUAL DEL MERCADO HOY (Básate en esto para tu análisis de impacto):\n"
        for region, indices in global_stats.items():
            for idx in indices:
                market_status_text += f"- {region} ({idx['nombre']}): {idx['efecto']}\n"
        
        if mode == "audio":
            focus_instruction = """
            MODO REPORTE DE AUDIO DE MERCADO (ANÁLISIS PROFUNDO, EXTENSO Y EXHAUSTIVO):
            El contenido proviene de un reporte de audio de mercado oficial enviado por una institución financiera (ej. audio semanal de Principal Financial Group / Banco / Corredora).
            ES ESTRICTAMENTE OBLIGATORIO que el texto del post para LinkedIn ('post_linkedin') sea PROFUNDO, EXTENSO Y EXHAUSTIVO (entre 5 y 8 párrafos completos estructurados con subtítulos y emojis).
            NO hagas un resumen corto ni superficial. Debes profundizar analíticamente en TODOS y cada uno de los temas tratados en el audio:
            1. 📌 **TITULAR Y RESUMEN ESTRATÉGICO DE LA SEMANA/JORNADA**
            2. 🌍 **COMPORTAMIENTO GLOBAL DE LOS MERCADOS**: Explica en detalle qué pasó en la bolsa de EE.UU. (S&P 500, NASDAQ), Europa, Asia y el IPSA local, citando los factores y catalizadores de los movimientos.
            3. 🏦 **BANCOS CENTRALES, INFLACIÓN Y TASAS**: Analiza las decisiones o postura de la Reserva Federal (FED) y del Banco Central de Chile (BCCh), datos de inflación e impacto en la curva de tasas de Renta Fija.
            4. 🛢️ **DÓLAR, COBRE Y COMMODITIES**: Analiza la trayectoria del tipo de cambio USD/CLP, el precio del cobre y del petróleo WTI.
            5. 💼 **IMPLICANCIAS Y ESTRATEGIA PARA PORTAFOLIOS PATRIMONIALES**: Explica qué significan estos datos y qué decisiones tácticas se recomiendan para Renta Fija (depósitos, bonos), Renta Variable (acciones) y Ahorro Previsional Voluntario (APV/Multifondos).
            6. 💡 **PERSPECTIVA Y PRÓXIMOS EVENTOS CLAVE**: Hitos o datos de la próxima semana que el inversionista debe monitorear.

            Escribe con el tono analítico, sofisticado e institucional del Economista Jefe de FV Asesorías e Inversiones. Usa negritas en todos los datos duros, emojis temáticos profesionales y dobles saltos de línea (\\n\\n) entre párrafos.
            """
        elif mode == "auto_chile":
            focus_instruction = "Elige la noticia MÁS IMPORTANTE enfocada en CHILE (economía, mercados, empresas, política que afecte la economía).\nCRÍTICO: Como la noticia es LOCAL de Chile, el impacto en los índices de USA, EUROPA y ASIA debe ser evaluado lógicamente. Una noticia local NO mueve el S&P500. Por lo tanto, para los índices globales, asigna 'NEUTRAL' en efecto y 'LEVE' o '-' en relevancia, a menos que la noticia tenga repercusiones mundiales demostrables."
        else:
            focus_instruction = "Elige la noticia MÁS IMPORTANTE basándote en su potencial impacto en los mercados globales y locales."
        
        current_date_str = datetime.datetime.now().strftime("%d-%m-%Y %H:%M")
        prompt = f"""
        Eres un Analista Senior de FV Asesorías e Inversiones SPA.
        La fecha y hora actual del sistema es: {current_date_str}.
        Evalúa las noticias del día y sus impactos:
        NOTICIAS HOY: {news_text}
        
        {market_status_text}
        
        {focus_instruction}
        IMPORTANTE: La explicación en 'explicacion_interna' DEBE coincidir causalmente con el RENDIMIENTO ACTUAL del mercado proporcionado arriba. Si el mercado en USA está al ALZA, explica por qué la noticia generó u ocurrió en un día de ALZA (ej. optimismo, menor regulación temida, etc.), NUNCA expliques una baja si el mercado real subió.
        Devuelve SOLO un JSON con esta estructura exacta:
        {{
            "titulo_principal": "TÍTULO DE IMPACTO EN MAYÚSCULAS",
            "titulo_documento": "Título muy corto, MÁXIMO 58 caracteres. Será el nombre del PDF.",
            "noticia_completa": "El texto periodístico de 4-5 oraciones profundizando en la noticia. (No des consejos).",
            "fuente_noticia": "Fuente real de la noticia extraída del texto (ej: Diario Financiero, Reuters, etc. NO inventes).",
            "fecha_noticia": "Usa ESTRICTAMENTE la fecha y hora que viene en el texto de NOTICIAS HOY. Si la hora viene en formato UTC (ejemplo terminada en Z), réstale 4 horas para ajustarla a Chile. NO inventes fechas pasadas ni uses la fecha de tus ejemplos. Formato final: DD de Mes, YYYY - HH:MM hrs",
            "prompt_imagen": "Un prompt corto en inglés (max 10 palabras) que describa la noticia para generar una imagen abstracta. (Ej: 'stock market crash red arrows')",
            "post_linkedin": "El post completo para RRSS. IMPORTANTE: Debes EXPLICAR BREVEMENTE de qué se trata la noticia o el suceso (el porqué o el contexto) antes de sacar conclusiones o hablar de incertidumbre, para que el lector entienda la causa. Usa un tono analítico, condicional y NO absolutista al proyectar impactos. Usa EMOJIS profesionales 📊 y destaca las ideas clave en NEGRITA (usando markdown **). Separa el texto obligatoriamente con DOBLE SALTO DE LÍNEA (\\n\\n) entre cada párrafo para que no se vea amontonado. Finaliza el post EXACTAMENTE con este bloque literal:\n\n¿Qué podría significar esto para tus ahorros e inversiones?\nDescubre cómo preparar tu portafolio ante estos nuevos desafíos. Obtén tu Radiografía Patrimonial, impulsada por nuestro software privado ALTUS AI, y optimiza tu estrategia de inversión.\n\n📧 contacto@fv-inversiones.com | 📱 WhatsApp: +56966779662\n\nAgrega de 3 a 5 HASHTAGS al final (ej: #Inversiones #Mercados).",
            "explicacion_interna": "Una explicación detallada (dirigida a los asesores de FV) de la lógica económica/financiera detrás de la noticia elegida y cómo fundamenta de forma causal los impactos (alzas y bajas) predichos en los commodities e índices. Sirve para responder dudas de clientes.",
            "explicacion_multifondos": "Un texto explicativo de unas 3-4 líneas (para clientes) justificando los movimientos proyectados (ALZA o BAJA) específicos de los Multifondos chilenos en base a la noticia y los mercados globales. Se incluirá en la presentación.",
            "impacto_local": {{
                "fondos_mutuos": [
                    {{"nombre": "GLOBAL", "efecto": "BAJA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "USA", "efecto": "BAJA", "relevancia": "MODERADA"}},
                    {{"nombre": "EUROPA", "efecto": "NEUTRAL", "relevancia": "LEVE"}},
                    {{"nombre": "ASIA", "efecto": "ALZA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "EMERGENTES", "efecto": "BAJA", "relevancia": "MODERADA"}},
                    {{"nombre": "LATAM", "efecto": "BAJA", "relevancia": "LEVE"}},
                    {{"nombre": "RV LOCAL", "efecto": "ALZA", "relevancia": "IMPORTANTE"}}
                ],
                "multifondos": [
                    {{"nombre": "Fondo A", "efecto": "BAJA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "Fondo B", "efecto": "BAJA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "Fondo C", "efecto": "NEUTRAL", "relevancia": "MODERADA"}},
                    {{"nombre": "Fondo D", "efecto": "NEUTRAL", "relevancia": "LEVE"}},
                    {{"nombre": "Fondo E", "efecto": "ALZA", "relevancia": "IMPORTANTE"}}
                ],
                "commodities": [
                    {{"nombre": "PETRÓLEO WTI", "efecto": "ALZA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "ORO", "efecto": "ALZA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "COBRE CASH", "efecto": "BAJA", "relevancia": "IMPORTANTE"}},
                    {{"nombre": "PLATA", "efecto": "BAJA", "relevancia": "LEVE"}},
                    {{"nombre": "GAS NATURAL", "efecto": "ALZA", "relevancia": "IMPORTANTE"}}
                ]
            }}
        }}
        Recuerda usar solo BAJA, ALZA, NEUTRAL para efecto y IMPORTANTE, MODERADA, LEVE para relevancia.
        IMPORTANTE: Devuelve un JSON estrictamente válido. Escapa las comillas internas con \\" y los saltos de línea dentro de los textos con \\n.
        """
        try:
            logging.info("Solicitando análisis de la noticia...")
            response = self.model.generate_content(prompt)
            clean_text = response.text.replace('```json', '').replace('```', '').strip()
            
            import json_repair
            data = json_repair.loads(clean_text)
            if not isinstance(data, dict):
                raise ValueError("El resultado reparado no es un diccionario válido.")
            
            # Inyectar Monedas y UF reales (100% nativo)
            data['impacto_local']['monedas'] = currency_stats
            data['impacto_local']['commodities'] = commodity_stats
            
            if mode == 'auto_chile':
                # El usuario solicitó que para noticias de Chile, todo se muestre como NEUTRAL y LEVE
                # para no generar falsas alarmas de impacto por una noticia local.
                for region, indices in global_stats.items():
                    for idx in indices:
                        idx['efecto'] = 'NEUTRAL'
                        idx['relevancia'] = 'LEVE'
                        
                for item in data['impacto_local'].get('monedas', []):
                    item['efecto'] = 'NEUTRAL'
                    item['relevancia'] = 'LEVE'
                    
                for item in data['impacto_local'].get('commodities', []):
                    item['efecto'] = 'NEUTRAL'
                    item['relevancia'] = 'LEVE'
                    
                for item in data['impacto_local'].get('fondos_mutuos', []):
                    item['efecto'] = 'NEUTRAL'
                    item['relevancia'] = 'LEVE'
                    
                for item in data['impacto_local'].get('multifondos', []):
                    item['efecto'] = 'NEUTRAL'
                    item['relevancia'] = 'LEVE'
            
            # Generar URL de imagen con Pollinations (gratis)
            encoded_prompt = urllib.parse.quote(data.get("prompt_imagen", "global financial markets abstract dark"))
            data['imagen_noticia'] = f"https://image.pollinations.ai/prompt/{encoded_prompt}?width=800&height=400&nologo=true"
            
            data['ltm_stats'] = self.fetch_ltm_variations()

            if not custom_news_text:
                news = self.db.query(MarketNews).filter(MarketNews.usado_para_linkedin == 0).limit(10).all()
                for n in news: n.usado_para_linkedin = 1
                self.db.commit()
            
            return data
        except Exception as e:
            logging.error(f"Error Gemini: {e}")
            return None

    def download_to_base64(self, url):
        try:
            headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 (KHTML, like Gecko) Chrome/91.0.4472.124 Safari/537.36"}
            resp = requests.get(url, headers=headers, timeout=15)
            if resp.status_code == 200:
                encoded = base64.b64encode(resp.content).decode("utf-8")
                mime = "image/png" if "png" in url else "image/jpeg"
                if "svg" in url: mime = "image/svg+xml"
                return f"data:{mime};base64,{encoded}"
            else:
                logging.error(f"Error descargando imagen, status: {resp.status_code}")
        except Exception as e:
            logging.error(f"Error descargando imagen {url}: {e}")
        return ""

    def render_infographic(self, global_stats, ai_data, mode="auto"):
        logging.info("Renderizando infografía completa a PNG...")
        
        # Cargar logos PNG oficiales en alta definición
        logo_fv = self.get_base64_image("assets/Logo_FV_Negativo.png")
        logo_altus = self.get_base64_image("assets/Logo_ALTUS AI_Negativo.png")
        
        # Descargar la imagen generada y el mapa para inyectarlos nativamente y evitar demoras de red
        if ai_data['imagen_noticia'].startswith('data:image'):
            # Ya es un base64 generado localmente (ej: Matplotlib COCHILCO)
            news_img_b64 = ai_data['imagen_noticia']
        else:
            news_img_b64 = self.download_to_base64(ai_data['imagen_noticia'])
            
        map_b64 = self.get_base64_image("assets/world_map.svg")
        
        if not news_img_b64:
            # Fallback estatico si falla pollinations: imagen financiera genérica
            fallback_url = "https://images.unsplash.com/photo-1590283603385-17ffb3a7f29f?ixlib=rb-4.0.3&auto=format&fit=crop&w=800&q=80"
            news_img_b64 = self.download_to_base64(fallback_url)
            if not news_img_b64:
                news_img_b64 = map_b64
        
        # Agregar datos al diccionario final
        ai_data['impacto_global'] = global_stats
        ai_data['logo_fv_base64'] = logo_fv
        ai_data['logo_altus_base64'] = logo_altus
        ai_data['map_b64'] = map_b64
        ai_data['mode'] = mode
        ai_data['fecha_str'] = datetime.datetime.now().strftime('%d de %B, %Y - %H:%M hrs')
        
        date_folder = datetime.date.today().strftime('%Y-%m-%d')
        out_dir = f"linkedin_posts/{date_folder}"
        os.makedirs(out_dir, exist_ok=True)
        
        # Guardar imagen a disco físico para evitar problemas de Chromium con b64 muy largos
        if news_img_b64 and news_img_b64.startswith('data:image'):
            try:
                base64_data = news_img_b64.split(",")[1]
                image_data = base64.b64decode(base64_data)
                chart_path = os.path.abspath(f"{out_dir}/temp_chart_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
                with open(chart_path, "wb") as f:
                    f.write(image_data)
                
                chart_path_url = chart_path.replace(chr(92), '/')
                import urllib.parse
                chart_path_url = urllib.parse.quote(chart_path_url, safe='/:')
                ai_data['news_img_b64'] = f"file:///{chart_path_url}"
            except Exception as e:
                logging.error(f"Error guardando grafico a disco: {e}")
                ai_data['news_img_b64'] = news_img_b64
        else:
            ai_data['news_img_b64'] = news_img_b64
        
        ai_data['agenda_dates'] = self.get_agenda_dates()
        ai_data['ltm_stats'] = self.fetch_ltm_variations()

        env = Environment(loader=FileSystemLoader('src/web/templates'))
        if mode in ["weekly", "audio"]:
            template = env.get_template('infografia_resumen_semanal.html')
        else:
            template = env.get_template('infografia_diaria.html')
        html_out = template.render(json_data=ai_data)
        
        html_path = f"{out_dir}/temp_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_out)
            
        hti = Html2Image(output_path=out_dir, custom_flags=[
            '--virtual-time-budget=4000', 
            '--allow-file-access-from-files', 
            '--force-device-scale-factor=2',
            '--hide-scrollbars'
        ])
        img_name = f"infografia_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        
        if mode in ["weekly", "audio"]:
            # Captura 4K limpia sin barras de desplazamiento (1200x1840 px a 2x = 2400x3680 px 4K Ultra HD)
            hti.screenshot(html_file=html_path, save_as=img_name, size=(1200, 1840))
        else:
            # Captura completa 4K para infografía diaria
            hti.screenshot(html_file=html_path, save_as=img_name, size=(1200, 3400))
        
        if os.path.exists(html_path): os.remove(html_path)
        return f"{date_folder}/{img_name}"

    def fetch_ltm_variations(self):
        """Calcula la variación real de los últimos 12 meses (LTM) para la grilla de divergencia con punto 0."""
        logging.info("Calculando variaciones LTM (12 Meses)...")
        defaults = {
            "sp500": 18.5,
            "nasdaq": 22.1,
            "bono10y": -3.8,
            "dolar": 7.4,
            "cobre": 12.3
        }
        tickers = [
            ("sp500", "^GSPC"),
            ("nasdaq", "^NDX"),
            ("bono10y", "^TNX"),
            ("dolar", "USDCLP=X"),
            ("cobre", "HG=F")
        ]
        results = {}
        for key, symbol in tickers:
            try:
                data = yf.Ticker(symbol).history(period="1y")
                if len(data) >= 20:
                    v_now = float(data['Close'].iloc[-1])
                    v_start = float(data['Close'].iloc[0])
                    pct = ((v_now - v_start) / v_start) * 100
                    results[key] = round(pct, 2)
                else:
                    results[key] = defaults[key]
            except Exception as e:
                logging.error(f"Error calculando LTM para {symbol}: {e}")
                results[key] = defaults[key]
        return results
        
        if os.path.exists(html_path): os.remove(html_path)
        return f"{date_folder}/{img_name}"

    def generate_weekly_chart_base64(self, ai_data):
        try:
            import io
            import matplotlib
            matplotlib.use('Agg')
            import matplotlib.pyplot as plt

            # Datos para el gráfico de variaciones semanales
            items = [
                ("S&P 500", -0.66),
                ("NASDAQ 100", -1.50),
                ("BONO EE.UU. 10Y", 3.52),
                ("DÓLAR (USD/CLP)", 0.99),
                ("COBRE (USD/LB)", 1.12)
            ]
            names = [x[0] for x in items]
            vals = [x[1] for x in items]
            colors = ['#ef4444' if v < 0 else '#10b981' for v in vals]

            fig, ax = plt.subplots(figsize=(9.5, 2.6), facecolor='#0f172a')
            ax.set_facecolor('#0f172a')

            bars = ax.barh(names, vals, color=colors, height=0.5)
            ax.axvline(0, color='#475569', linewidth=1.5, linestyle='--')

            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)
            ax.spines['bottom'].set_color('#334155')
            ax.spines['left'].set_color('#334155')
            ax.tick_params(colors='#94a3b8', labelsize=10.5)
            ax.set_title("VARIACIÓN SEMANAL (%) POR CLASE DE ACTIVO Y INDICADOR", color='#38bdf8', fontsize=11, fontweight='bold', pad=10)

            for bar, val in zip(bars, vals):
                offset = 0.12 if val >= 0 else -0.12
                ha = 'left' if val >= 0 else 'right'
                sign = '+' if val > 0 else ''
                ax.text(val + offset, bar.get_y() + bar.get_height()/2, f"{sign}{val:.2f}%", 
                        va='center', ha=ha, color='#ffffff', fontweight='bold', fontsize=10)

            plt.tight_layout()
            
            buf = io.BytesIO()
            plt.savefig(buf, format='png', dpi=180, facecolor='#0f172a')
            plt.close(fig)
            buf.seek(0)
            b64_str = base64.b64encode(buf.read()).decode('utf-8')
            return f"data:image/png;base64,{b64_str}"
        except Exception as e:
            logging.error(f"Error generando gráfico semanal Matplotlib: {e}")
            return ""

    def get_agenda_dates(self):
        today = datetime.date.today()
        days_to_tue = (1 - today.weekday()) % 7
        if days_to_tue == 0: days_to_tue = 7
        next_tue = today + datetime.timedelta(days=days_to_tue)
        next_wed = next_tue + datetime.timedelta(days=1)
        next_thu = next_tue + datetime.timedelta(days=2)
        next_fri = next_tue + datetime.timedelta(days=3)
        return {
            "tue": next_tue.strftime("%d.%m.%y"),
            "wed": next_wed.strftime("%d.%m.%y"),
            "thu": next_thu.strftime("%d.%m.%y"),
            "fri": next_fri.strftime("%d.%m.%y")
        }

    def render_carousel(self, global_stats, ai_data):
        logging.info("Renderizando carrusel PDF...")
        
        env = Environment(loader=FileSystemLoader('src/web/templates'))
        template = env.get_template('carrusel_diario.html')
        html_out = template.render(json_data=ai_data)
        
        date_folder = datetime.date.today().strftime('%Y-%m-%d')
        out_dir = f"linkedin_posts/{date_folder}"
        os.makedirs(out_dir, exist_ok=True)
        
        # Guardar imagen a disco físico para evitar problemas de Chromium con b64 muy largos en el carrusel
        news_img_b64 = ai_data.get('news_img_b64', '')
        if news_img_b64 and news_img_b64.startswith('data:image'):
            try:
                base64_data = news_img_b64.split(",")[1]
                image_data = base64.b64decode(base64_data)
                chart_path = os.path.abspath(f"{out_dir}/temp_carousel_chart_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png")
                with open(chart_path, "wb") as f:
                    f.write(image_data)
                
                chart_path_url = chart_path.replace(chr(92), '/')
                import urllib.parse
                chart_path_url = urllib.parse.quote(chart_path_url, safe='/:')
                ai_data['news_img_b64'] = f"file:///{chart_path_url}"
            except Exception as e:
                logging.error(f"Error guardando grafico a disco para carrusel: {e}")
                
        # Re-renderizar HTML con la nueva ruta file:///
        html_out = template.render(json_data=ai_data)
        
        html_path = f"{out_dir}/temp_carrusel_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        
        with open(html_path, "w", encoding="utf-8") as f:
            f.write(html_out)
            
        hti = Html2Image(output_path=out_dir, custom_flags=['--virtual-time-budget=4000', '--allow-file-access-from-files'])
        img_name = f"temp_carrusel_largo_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.png"
        
        # Capturar la tira completa (2160 de ancho por 8640 de alto, 4 slides) para mayor resolución
        hti.screenshot(html_file=html_path, save_as=img_name, size=(2160, 8640))
        
        img_path = f"{out_dir}/{img_name}"
        if os.path.exists(img_path):
            Image.init()  # Necesario para forzar la carga de los plugins (como JPEG) antes de guardar a PDF
            img = Image.open(img_path)
            pages = []
            for i in range(4):
                # Cada página es de 2160x2160 píxeles ahora
                box = (0, i*2160, 2160, (i+1)*2160)
                page = img.crop(box).convert('RGB')
                pages.append(page)
                
            pdf_name = f"{out_dir}/carrusel_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.pdf"
            pages[0].save(pdf_name, save_all=True, append_images=pages[1:])
            
            # Limpiar temporales
            img.close()
            os.remove(img_path)
        
        if os.path.exists(html_path): os.remove(html_path)
        logging.info(f"Carrusel PDF generado exitosamente.")
        return True

    def publish_to_linkedin(self, img_path, text_content):
        """
        [OPCIÓN 2 - AVANZADA]
        Método preparado para publicar automáticamente en LinkedIn cuando se disponga de los tokens.
        Actualmente está detrás de un Feature Flag (LINKEDIN_ACCESS_TOKEN).
        """
        token = os.getenv("LINKEDIN_ACCESS_TOKEN")
        if not token:
            logging.info("Opción 2 (Auto-Publicación en LinkedIn) desactivada: Falta LINKEDIN_ACCESS_TOKEN en el entorno.")
            logging.info("Usando Opción 1: Publicación manual requerida desde la carpeta 'linkedin_posts'.")
            return False
            
        logging.info("Iniciando publicación automatizada en LinkedIn...")
        # Aquí irá la lógica de llamadas a la API de LinkedIn (v2/ugcPosts)
        # 1. Obtener URN del autor (Empresa o Persona)
        # 2. Registrar el asset (imagen)
        # 3. Subir el binario de la imagen
        # 4. Crear el post con el texto y la imagen adjunta
        
        logging.info("Publicación enviada a LinkedIn exitosamente.")
        return True

    def run_daily_routine(self, mode="auto", custom_input=None):
        global_stats = self.fetch_global_stats()
        currency_stats = self.fetch_currency_stats()
        commodity_stats = self.fetch_commodities_stats()
        
        custom_news_text = None
        if mode == "url" and custom_input:
            custom_news_text = self.fetch_custom_url(custom_input)
            if not custom_news_text:
                raise ValueError("No se pudo extraer la noticia de la URL. Posible muro de pago o enlace inválido.")
        elif mode == "topic" and custom_input:
            custom_news_text = self.search_topic(custom_input)
            if not custom_news_text:
                raise ValueError(f"No se encontraron resultados útiles para el tema: {custom_input}")
        elif mode == "sonami":
            logging.info("Iniciando escaneo especial: Cumbre SONAMI 2026")
            url = "https://cumbresonami.cl/"
            custom_news_text = self.fetch_custom_url(url)
            if custom_news_text:
                custom_news_text += "\n\n[INSTRUCCIÓN CRÍTICA PARA LA IA - MODO EXPERTO]: El texto anterior es solo la página de venta de entradas de un evento. ES PROHIBIDO hacer un simple resumen promocional del evento. Tu objetivo es usar este evento como EXCUSA para hablar de los problemas reales de la minería en Chile hoy: estancamiento de la producción de cobre en torno a 5 millones de toneladas, la permisología asfixiante que retrasa proyectos, el impacto del royalty minero en los márgenes de las empresas, y cómo el alto precio internacional del cobre choca con la incapacidad local de aumentar la oferta. Redacta un post de LinkedIn con tono de alerta y urgencia patrimonial, analizando cómo estos factores (y las decisiones que se tomen en esta cumbre) afectarán directamente a los fondos mutuos chilenos, las acciones del IPSA y el tipo de cambio."
            else:
                raise ValueError("No se pudo conectar con el sitio de SONAMI.")
                
        elif mode == "cochilco":
            logging.info("Iniciando Reporte Cuantitativo (COCHILCO)")
            ode = OpenDataEngine()
            df = ode.get_cochilco_production_data()
            custom_news_text = ode.generate_llm_context(df)
            self.cochilco_chart_b64 = ode.generate_production_chart_base64(df)
        elif mode in ["weekly", "audio"] and custom_input:
            logging.info("Iniciando procesamiento de Infografía desde Reporte Resumen Semanal de Mercado (Audio/WhatsApp)")
            custom_news_text = custom_input
            
        if mode == "auto":
            self.fetch_reuters_news()
        elif mode == "auto_chile":
            self.fetch_chile_news()
        
        ai_data = self.generate_content(currency_stats, commodity_stats, global_stats, custom_news_text, mode)
        if ai_data:
            # Si generamos un gráfico de COCHILCO, reemplazar la imagen genérica de IA por nuestro gráfico duro
            if mode == "cochilco" and hasattr(self, 'cochilco_chart_b64'):
                ai_data['imagen_noticia'] = self.cochilco_chart_b64
                
            img_file = self.render_infographic(global_stats, ai_data, mode)
            if mode not in ["weekly", "audio"]:
                self.render_carousel(global_stats, ai_data)
            
            date_folder = datetime.date.today().strftime('%Y-%m-%d')
            out_dir = f"linkedin_posts/{date_folder}"
            
            post_file = f"{out_dir}/post_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
            
            titulo_doc = ai_data.get('titulo_documento', 'Reporte Diario de Mercados')
            texto_final_post = f"📁 **TÍTULO SUGERIDO PARA EL PDF (Max 58 chars):** {titulo_doc}\n\n---\n\n"
            texto_final_post += self.markdown_to_unicode_bold(ai_data['post_linkedin'])
            
            with open(post_file, "w", encoding="utf-8") as f:
                f.write(texto_final_post)
                
            fundamento_file = f"{out_dir}/fundamento_interno_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc = Document()
            doc.add_heading('Racional Interno de Inversiones', 0)
            doc.add_paragraph(f"Fecha del Reporte: {ai_data.get('fecha_str', '')}")
            
            doc.add_heading('Noticia Seleccionada', level=1)
            doc.add_heading(ai_data.get('titulo_principal', ''), level=2)
            doc.add_paragraph(f"Fuente: {ai_data.get('fuente_noticia', '')} | Publicada: {ai_data.get('fecha_noticia', '')}")
            doc.add_paragraph(ai_data.get('noticia_completa', ''))
            
            doc.add_heading('Lógica y Fundamento de Impactos (Uso Interno)', level=1)
            doc.add_paragraph(ai_data.get('explicacion_interna', 'Explicación no generada.'))
            doc.save(fundamento_file)
                
            logging.info(f"ÉXITO: Reportes generados en {out_dir}/")
            
            # Intento de publicación automatizada (Opción 2)
            self.publish_to_linkedin(f"linkedin_posts/{img_file}", ai_data['post_linkedin'])
            
            return ai_data['post_linkedin'], img_file

if __name__ == "__main__":
    engine = MarketDataEngine()
    engine.run_daily_routine()

import os
from datetime import datetime
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def generate_macro_docx(cliente_nombre: str, contenido_markdown: str, output_path: str):
    """
    Genera un informe institucional ejecutable en formato Word (.docx)
    con encabezados corporativos, disclaimers, Sobre FV Asesorías y ciberseguridad.
    """
    doc = Document()
    
    # Configuración de márgenes
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Encabezado con Logo Oficial FV desde la carpeta de marca assets/brand
    root_dir = os.path.dirname(os.path.dirname(os.path.dirname(__file__)))
    fv_logo_path = os.path.join(root_dir, "assets", "brand", "fv_logo_principal_light.png")
    if not os.path.exists(fv_logo_path):
        fv_logo_path = os.path.join(root_dir, "assets", "brand", "fv_logo_principal_trimmed.png")
    if os.path.exists(fv_logo_path):
        try:
            p_logo = doc.add_paragraph()
            p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run_logo = p_logo.add_run()
            run_logo.add_picture(fv_logo_path, width=Inches(1.8))
        except Exception:
            pass

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run_sub = title_p.add_run("DIGITAL FAMILY OFFICE ANALYTICS | FV ASESORÍAS & ALTUS AI\n")
    run_sub.font.size = Pt(8.5)
    run_sub.font.bold = True
    run_sub.font.color.rgb = RGBColor(10, 35, 66)
    
    run_date = title_p.add_run(f"Fecha de Emisión: {datetime.now().strftime('%d-%m-%Y')}\n")
    run_date.font.size = Pt(8)
    run_date.font.color.rgb = RGBColor(100, 116, 139)

    # Tabla de Datos del Reporte
    tbl = doc.add_table(rows=2, cols=2)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl.autofit = True
    
    cell_a = tbl.cell(0, 0)
    cell_a.text = f"Atención a: {cliente_nombre}"
    cell_b = tbl.cell(0, 1)
    cell_b.text = f"Fecha: {datetime.now().strftime('%d/%m/%Y')}"
    cell_c = tbl.cell(1, 0)
    cell_c.text = "Unidad: Inteligencia de Mercados & Asignación de Activos"
    cell_d = tbl.cell(1, 1)
    cell_d.text = "Certificado por: Altus AI Macro Engine"
    
    doc.add_paragraph()

    # Título Principal del Reporte
    h1 = doc.add_heading("Consenso Definitivo e Inteligencia de Mercado", level=1)
    h1.runs[0].font.color.rgb = RGBColor(16, 75, 60)
    h1.runs[0].font.size = Pt(18)

    # Procesar líneas de Markdown a párrafos Word
    lines = contenido_markdown.split('\n')
    for line in lines:
        line_s = line.strip()
        if not line_s:
            continue
        if line_s.startswith("# "):
            h = doc.add_heading(line_s[2:], level=1)
            h.runs[0].font.color.rgb = RGBColor(10, 35, 66)
        elif line_s.startswith("## "):
            h = doc.add_heading(line_s[3:], level=2)
            h.runs[0].font.color.rgb = RGBColor(16, 75, 60)
        elif line_s.startswith("### "):
            h = doc.add_heading(line_s[4:], level=3)
            h.runs[0].font.color.rgb = RGBColor(30, 41, 59)
        elif line_s.startswith("- ") or line_s.startswith("* "):
            p = doc.add_paragraph(line_s[2:], style='List Bullet')
            p.paragraph_format.space_after = Pt(4)
        else:
            p = doc.add_paragraph(line_s)
            p.paragraph_format.space_after = Pt(6)

    doc.add_paragraph()

    # Aviso Legal / Disclaimer
    p_disc = doc.add_paragraph()
    p_disc.paragraph_format.space_before = Pt(14)
    run_disc_t = p_disc.add_run("Aviso Legal: ")
    run_disc_t.bold = True
    run_disc_t.font.size = Pt(8.5)
    run_disc_t.font.color.rgb = RGBColor(100, 116, 139)
    run_disc_b = p_disc.add_run(
        "Las visiones y proyecciones macroeconómicas presentadas en este documento han sido procesadas mediante inteligencia artificial (Altus AI) cruzando múltiples visiones institucionales. Este documento no constituye una recomendación de inversión vinculante, sino una herramienta de información estratégica. Los mercados son volátiles y las rentabilidades pasadas no garantizan retornos futuros. FV Asesorías e Inversiones limita su responsabilidad al análisis cuantitativo."
    )
    run_disc_b.font.size = Pt(8.5)
    run_disc_b.font.color.rgb = RGBColor(100, 116, 139)

    # Sobre FV Asesorías e Inversiones
    p_fv = doc.add_paragraph()
    p_fv.paragraph_format.space_before = Pt(10)
    run_fv_t = p_fv.add_run("Sobre FV Asesorías e Inversiones:\n")
    run_fv_t.bold = True
    run_fv_t.font.size = Pt(9)
    run_fv_t.font.color.rgb = RGBColor(212, 175, 55) # Altus Gold
    run_fv_b = p_fv.add_run(
        "FV Asesorías e Inversiones somos un Multi-Family Office Digital impulsado por nuestro software cuantitativo privado de Inteligencia Artificial (ALTUS AI). Combinamos la agilidad tecnológica de una WealthTech con la exclusividad de una oficina patrimonial privada, auditando en 360° la situación tributaria, inmobiliaria, composición familiar, seguros e inversiones para proteger su legado a través de las generaciones."
    )
    run_fv_b.font.size = Pt(9)
    run_fv_b.font.color.rgb = RGBColor(55, 65, 81)

    # Ciberseguridad & Resguardo Patrimonial
    p_sec = doc.add_paragraph()
    p_sec.paragraph_format.space_before = Pt(10)
    run_sec_t = p_sec.add_run("🔒 Ciberseguridad & Resguardo Patrimonial: ")
    run_sec_t.bold = True
    run_sec_t.font.size = Pt(8.5)
    run_sec_t.font.color.rgb = RGBColor(10, 35, 66)
    run_sec_b = p_sec.add_run(
        "Toda la información analizada por Altus AI se encuentra protegida bajo cifrado nativo AES-256 bits y transmisión TLS 1.3 de grado bancario. Garantizamos estricta confidencialidad bajo Secreto Patrimonial y cumplimiento riguroso de la Ley N° 19.628 de Protección de Datos Personales en Chile."
    )
    run_sec_b.font.size = Pt(8.5)
    run_sec_b.font.color.rgb = RGBColor(51, 65, 85)

    doc.save(output_path)
    return output_path

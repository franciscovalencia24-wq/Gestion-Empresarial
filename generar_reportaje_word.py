import os
try:
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    import sys
    import subprocess
    subprocess.check_call([sys.executable, '-m', 'pip', 'install', 'python-docx'])
    from docx import Document
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

def crear_documento():
    doc = Document()
    
    # Titulo
    title = doc.add_heading('De la Asesoría Tradicional al Primer Digital Family Office de Chile', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('Evolución, Tecnología y el Nuevo Rostro de la Gestión Patrimonial')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].bold = True
    
    doc.add_paragraph()
    
    p1 = doc.add_paragraph('Por años, la industria financiera en Chile ha operado bajo las mismas reglas. El cliente, cansado y confundido, deposita su dinero en un banco, contrata seguros en una corredora y confía sus impuestos a un contador. Sin embargo, nadie tiene la "foto completa". En este modelo fragmentado, el ejecutivo bancario tradicional suele actuar más como un vendedor presionado por metas mensuales que como un verdadero guardián de los intereses de la familia.')
    
    p2 = doc.add_paragraph('Frente a este escenario de incertidumbre, FV Asesorías e Inversiones decidió que era el momento de romper el molde. Nos dimos cuenta de que las familias no necesitan que les vendan más "productos rentables"; necesitan paz mental, certeza estadística y un blindaje real frente a las crisis.')
    
    doc.add_heading('El Salto Cuántico: Inteligencia Artificial al Servicio del Patrimonio', level=1)
    
    doc.add_paragraph('Para lograr este nivel de protección, construimos algo inédito en el mercado nacional: un ecosistema tecnológico impulsado por Inteligencia Artificial y modelos estadísticos institucionales. Nuestra plataforma es hoy un cerebro digital capaz de hacer en segundos lo que a un equipo de analistas de banca privada le tomaría semanas de trabajo manual.')
    
    doc.add_paragraph('¿Qué significa esto en la práctica? Significa que nuestro sistema puede "leer" automáticamente los documentos bancarios de un cliente, consolidando su patrimonio en un abrir y cerrar de ojos. Significa que, utilizando tecnología de Recuperación Aumentada (RAG), nuestro motor conoce las sutilezas de la Ley de Herencia chilena y propone vehículos de inversión libres de impuestos. Y lo más importante: significa que sometemos el flujo de caja del cliente a simulaciones extremas (Monte Carlo), garantizándole matemáticamente hasta qué edad le durará su capital manteniendo su estilo de vida.')
    
    quote = doc.add_paragraph('"La Inteligencia Artificial no reemplaza nuestra cercanía; nos otorga superpoderes para ser los arquitectos definitivos del patrimonio de nuestras familias."')
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.runs[0].italic = True
    
    doc.add_heading('El Nuevo Símbolo de una Nueva Era', level=1)
    doc.add_paragraph('Todo este salto tecnológico debía reflejarse en nuestra identidad. El antiguo logo de FV Asesorías (la moneda) representaba con orgullo nuestro pasado: la figura del Asesor Financiero tradicional que te ayudaba a guardar y proteger tu dinero.')
    
    doc.add_paragraph('Pero hoy somos mucho más que eso.')
    
    doc.add_paragraph('Nuestro nuevo logo es el símbolo del futuro. Representa la figura del Arquitecto Tecnológico, aquel que estructura, blinda y proyecta tu patrimonio global sin sesgos comerciales. Sus líneas modernas y colores sobrios reflejan la exactitud matemática y la sofisticación de un verdadero Digital Family Office.')
    
    doc.add_heading('Un Futuro Exclusivo y Cupos Especiales', level=1)
    doc.add_paragraph('Hoy, nuestro modelo es una realidad operativa. Sin embargo, nuestro enfoque no es la masividad, sino la excelencia absoluta. Por ello, hemos decidido lanzar esta tecnología a través de un programa piloto exclusivo, atendiendo solo a un grupo rigurosamente seleccionado de familias para garantizar un nivel de servicio "CEO-style" inigualable.')
    
    # CTA Final
    cta = doc.add_paragraph()
    cta_run = cta.add_run('Invitación Especial para Lectores: ')
    cta_run.bold = True
    cta.add_run('Hemos reservado un número muy limitado de cupos en este programa piloto de forma exclusiva para los lectores de esta revista. Si desea experimentar el futuro de la gestión patrimonial y someter su portafolio actual a un diagnóstico bajo nuestra tecnología de Inteligencia Artificial, contáctenos directamente al correo electrónico ')
    
    email = cta.add_run('contacto@fv-inversiones.com')
    email.bold = True
    email.underline = True
    
    cta.add_run(' o comuníquese a nuestro teléfono ')
    
    fono = cta.add_run('+56 9 1234 5678') # Placeholder a cambiar si es necesario
    fono.bold = True
    
    cta.add_run('. Dé el primer paso hacia su paz mental.')
    
    doc.save('Reportaje_Revista_FV.docx')
    print("Reportaje guardado como Reportaje_Revista_FV.docx")

if __name__ == '__main__':
    crear_documento()
